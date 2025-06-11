import io
import os
import PyPDF2
import requests

from dotenv import load_dotenv
from django.conf import settings
from langchain.llms import OpenAI
from docx import Document as DocxDocument
from langchain.vectorstores import Chroma
from langchain.docstore.document import Document
from langchain.chains import RetrievalQAWithSourcesChain
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter


load_dotenv()
OPENAI_API_KEY = os.environ["OPENAI_API_KEY"]


class LLM:
    def __init__(self, db_path):
        self.embeddings = OpenAIEmbeddings()
        self.db_path = db_path + "/vectorDB"

    def load_data(self, url_dict):
        files = [file.asset.url for file in url_dict]
        return files

    def create_docs(self, file_path):
        total_docs = []
        files = self.load_data(file_path)
        for file in files:
            response = requests.get(file)
            if response.status_code != 200:
                raise ValueError(
                    f"Failed to fetch file from source. Status code: {response.status_code}"
                )

            file_extension = file.split(".")[-1].lower()
            text = ""

            if file_extension == "pdf":
                bytes_data = io.BytesIO(response.content)
                pdf_reader = PyPDF2.PdfReader(bytes_data)
                text = "".join([page.extract_text() for page in pdf_reader.pages])

            elif file_extension == "docx":
                bytes_data = io.BytesIO(response.content)
                docx_doc = DocxDocument(bytes_data)
                text = "\n".join([paragraph.text for paragraph in docx_doc.paragraphs])

            elif file_extension == "txt":
                text = response.content.decode("utf-8")

            doc = [Document(page_content=text, metadata={"source": file})]

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000, chunk_overlap=200, length_function=len
            )
            docs = text_splitter.split_documents(doc)
            total_docs.extend(docs)

        return total_docs

    # stores the embeddings to a db
    def store_embeddings(self, docs):
        try:
            # Chroma
            db = Chroma.from_documents(
                docs, self.embeddings, persist_directory=self.db_path
            )
            db.persist()
            return True
        except Exception as e:
            print("Error storing embeddings:", e)
            return False

    def getQA(self):
        # Chroma
        vectordb = Chroma(
            persist_directory=self.db_path,
            embedding_function=OpenAIEmbeddings(),
        )
        retriever = vectordb.as_retriever(search_kwargs={"k": 3})

        qa_chain = RetrievalQAWithSourcesChain.from_chain_type(
            llm=OpenAI(),
            chain_type="stuff",
            retriever=retriever,
            return_source_documents=False,
        )

        return qa_chain


class Classifier:
    def __init__(self, db_path):
        self.llm = LLM(db_path)

    def train(self, doc_path):
        docs = self.llm.create_docs(doc_path)
        self.llm.store_embeddings(docs)

    def convert_result_to_dict(self, result):
        if isinstance(result, dict):
            return {
                key: self.convert_result_to_dict(value) for key, value in result.items()
            }
        elif isinstance(result, list):
            return [self.convert_result_to_dict(item) for item in result]
        elif isinstance(result, Document):
            # Remove "\n" characters from page_content and metadata
            cleaned_page_content = result.page_content.replace("\n", "")
            cleaned_metadata = {
                key: value.replace("\n", "") if isinstance(value, str) else value
                for key, value in result.metadata.items()
            }
            return {
                "page_content": cleaned_page_content,
                "metadata": cleaned_metadata,
            }
        elif isinstance(result, str):
            # Split the "sources" field by ', ' and remove any whitespace around URLs
            if result.startswith("[") and result.endswith("]"):
                # Remove square brackets from the "sources" field if they exist
                result = result[1:-1]
            sources_list = [url.strip() for url in result.split(",")]
            return sources_list
        return result

    def ask(self, question):
        qa = self.llm.getQA()
        response = qa({"question": question})
        sources = response.get("sources", "")

        # Convert Document objects into dictionaries recursively within the sources list
        if isinstance(sources, list):
            response["sources"] = self.convert_sources_to_dict(sources)

        # Convert the entire response dictionary into dictionaries recursively
        response = self.convert_result_to_dict(response)

        # Concatenate the answer strings into a single string without newlines
        if "answer" in response:
            response["answer"] = " ".join(response["answer"])

        return response
